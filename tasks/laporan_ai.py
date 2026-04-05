#!/usr/bin/env python3
import sys
import os
import subprocess
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TOKEN = "8187479612:AAE3ZF-qItlSd2MhoZX2e0NbvTBaeAVemhc"
CHAT_ID = "7250558059"

def send_file(filepath):
    filename = os.path.basename(filepath)
    cmd = [
        "curl", "-s", "-F", f"chat_id={CHAT_ID}",
        "-F", f"document=@{filepath}",
        "-F", f"caption=Laporan Lengkap AI",
        f"https://api.telegram.org/bot{TOKEN}/sendDocument"
    ]
    subprocess.run(cmd, capture_output=True)

def buat_laporan():
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    
    def add_para(text, bold=False, center=False):
        p = doc.add_paragraph(text)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold:
            for run in p.runs:
                run.bold = True
        return p
    
    # COVER
    doc.add_paragraph()
    add_para("LAPORAN", True, True)
    add_para("INTELIGENCE ARTIFICIAL (AI)", True, True)
    add_para("DALAM KEHIDUPAN MANUSIA", True, True)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    for label, nilai in [("Nama", "Ahdi Khalida Fathir"), ("Kelas", "X IPA"), ("Mapel", "Informatika"), ("Tanggal", "April 2026")]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(nilai)
    
    doc.add_page_break()
    
    # BAB I
    h = doc.add_heading("BAB I. PENDAHULUAN", level=1)
    h.runs[0].font.size = Pt(14)
    add_para("Kecerdasan Buatan atau Artificial Intelligence (AI) merupakan cabang ilmu komputer yang berfokus pada pembuatan mesin yang dapat melakukan tugas-tugas membutuhkan kecerdasan manusia. Sejak diperkenalkan tahun 1956 di Dartmouth Conference, AI telah berkembang sangat signifikan.")
    add_para("Di era digital saat ini, AI telah merambah berbagai aspek kehidupan manusia. Dari asisten virtual seperti Siri dan Alexa, sistem rekomendasi di platform streaming seperti Netflix, hingga mobil self-driving yang dikembangkan oleh Tesla dan Waymo. Perkembangan ini mengubah cara manusia bekerja dan berinteraksi dengan teknologi.")
    add_para("Laporan ini disusun untuk memberikan pemahaman komprehensif mengenai AI, mulai dari definisi, sejarah, jenis-jenis, cara kerja, hingga dampaknya terhadap berbagai aspek kehidupan manusia.")
    
    # BAB II
    h = doc.add_heading("BAB II. TUJUAN", level=1)
    h.runs[0].font.size = Pt(14)
    add_para("1. Memahami pengertian dan konsep dasar Artificial Intelligence (AI)")
    add_para("2. Mengetahui sejarah dan perkembangan AI dari masa ke masa")
    add_para("3. Mengidentifikasi jenis-jenis AI berdasarkan kapabilitas dan fungsinya")
    add_para("4. Memahami cara kerja dan teknologi yang mendasari AI")
    add_para("5. Menganalisis implementasi AI dalam berbagai bidang kehidupan")
    add_para("6. Mengidentifikasi manfaat dan dampak negatif AI")
    add_para("7. Memahami prospek dan tantangan AI di masa depan")
    
    # BAB III
    h = doc.add_heading("BAB III. LANDASAN TEORI", level=1)
    h.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("3.1 Pengertian Artificial Intelligence").bold = True
    add_para("Artificial Intelligence (AI) atau Kecerdasan Buatan adalah cabang ilmu komputer yang berfokus pada pembuatan mesin yang dapat melakukan tugas-tugas yang membutuhkan kecerdasan manusia. Menurut Stuart Russell dan Peter Norvig, AI didefinisikan sebagai studi tentang agen yang menerima persepsi dari lingkungan dan execute actions.")
    add_para("John McCarthy mendefinisikan AI sebagai science and engineering of making intelligent machines, especially intelligent computer programs.")
    
    p = doc.add_paragraph()
    p.add_run("3.2 Jenis-jenis AI").bold = True
    add_para("Berdasarkan kapabilitasnya:")
    add_para("- Narrow AI (Weak AI): AI untuk tugas spesifik tertentu seperti Siri, Alexa, chatbot")
    add_para("- General AI (Strong AI): AI dengan kemampuan kognitif seperti manusia (konsep teoretis)")
    add_para("- Super AI: AI yang melampaui kecerdasan manusia (masih spekulasi)")
    add_para("Berdasarkan fungsinya:")
    add_para("- Reactive Machines: merespons stimulus tanpa memori (contoh: IBM Deep Blue)")
    add_para("- Limited Memory: menggunakan pengalaman masa lalu (contoh: Self-driving cars)")
    add_para("- Theory of Mind: memahami emosi dan niat manusia (dalam pengembangan)")
    add_para("- Self-Aware AI: memiliki kesadaran diri (belum ada)")
    
    p = doc.add_paragraph()
    p.add_run("3.3 Teknologi Pendukung AI").bold = True
    add_para("- Machine Learning: subbidang AI yang memungkinkan mesin belajar dari data")
    add_para("- Deep Learning: subbidang ML menggunakan neural networks berlapis")
    add_para("- NLP (Natural Language Processing): kemampuan AI memahami bahasa manusia")
    add_para("- Computer Vision: kemampuan AI menginterpretasi gambar dan video")
    add_para("- Robotics: integrasi AI dengan sistem mekanik")
    
    # BAB IV
    h = doc.add_heading("BAB IV. METODE", level=1)
    h.runs[0].font.size = Pt(14)
    add_para("Metode yang digunakan:")
    add_para("1. Studi Literatur - Pengumpulan data dari buku, jurnal ilmiah, dan artikel terpercaya. Penelitian terhadap berbagai sumber akademik tentang AI.")
    add_para("2. Observasi - Pengamatan terhadap implementasi AI dalam kehidupan sehari-hari dan identifikasi contoh-contoh aplikasi AI.")
    add_para("3. Analisis Data - Pengolahan informasi dari berbagai sumber dan sintesis data untuk mendapatkan kesimpulan.")
    
    # BAB V
    h = doc.add_heading("BAB V. DATA PENGAMATAN", level=1)
    h.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("5.1 Implementasi AI dalam Berbagai Bidang").bold = True
    
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Bidang", "Contoh Aplikasi AI", "Keterangan"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    data = [
        ("Kesehatan", "Diagnosis medis, robot bedah", "Meningkatkan akurasi"),
        ("Pendidikan", "Adaptive learning", "Personalisasi belajar"),
        ("Ekonomi", "Algoritma trading", "Keputusan cepat"),
        ("Transportasi", "Mobil otonom", "Keselamatan naik"),
        ("Hiburan", "Rekomendasi konten", "Pengalaman lebih baik"),
    ]
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("5.2 Perkembangan Pengguna AI").bold = True
    add_para("- 2020: 1,5 miliar pengguna AI")
    add_para("- 2023: 3 miliar pengguna AI")
    add_para("- 2025: Diperkirakan 4,5 miliar")
    add_para("- 2030: Diperkirakan 6,5 miliar")
    
    # BAB VI
    h = doc.add_heading("BAB VI. PEMBAHASAN", level=1)
    h.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("6.1 Manfaat AI").bold = True
    add_para("a) Efisiensi: AI menyelesaikan tugas lebih cepat dan akurat. Robot AI dapat beroperasi 24 jam dengan kesalahan minimal.")
    add_para("b) Pengambilan Keputusan: AI menganalisis data besar membantu keputusan lebih akurat.")
    add_para("c) Inovasi Kesehatan: AI mendeteksi kanker 95% akurasi, dikembangkan obat baru lebih cepat.")
    add_para("d) Personalisasi: AI membuat pengalaman digital lebih personal dan relevan.")
    
    p = doc.add_paragraph()
    p.add_run("6.2 Dampak Negatif AI").bold = True
    add_para("a) Pengangguran: AI diprediksi替换85 juta pekerjaan pada 2025.")
    add_para("b) Bias Algoritma: AI dapat mewarisi bias dari data pelatihan.")
    add_para("c) Privasi: Penggunaan AI dalam pengawasan mengancam privasi.")
    add_para("d) Deepfake: AI digunakan membuat konten palsu yang meyakinkan.")
    
    p = doc.add_paragraph()
    p.add_run("6.3 Etika AI").bold = True
    add_para("- Transparansi: Keputusan AI harus dapat dijelaskan")
    add_para("- Keadilan: AI tidak boleh diskriminatif")
    add_para("- Privasi: Data pengguna harus dilindungi")
    add_para("- Keamanan: Sistem AI harus aman")
    add_para("- Akuntabilitas: Ada pihak bertanggung jawab")
    
    # BAB VII
    h = doc.add_heading("BAB VII. KESIMPULAN", level=1)
    h.runs[0].font.size = Pt(14)
    add_para("1. Artificial Intelligence merupakan teknologi yang memungkinkan mesin meniru kecerdasan manusia. AI telah berkembang pesat sejak 1956.")
    add_para("2. AI memiliki berbagai jenis dengan kapabilitas berbeda, dari Narrow AI hingga konsep Super AI.")
    add_para("3. Implementasi AI telah merambah bidang kesehatan, pendidikan, ekonomi, transportasi, dan hiburan dengan manfaat signifikan.")
    add_para("4. AI membawa dampak negatif seperti pengangguran, bias algoritma, dan ancaman privasi. Pengembangan AI harus diiringi regulasi dan etika ketat.")
    add_para("5. Masa depan AI sangat menjanjikan dengan prediksi mayoritas populasi dunia akan menggunakan AI pada 2030.")
    
    # BAB VIII
    h = doc.add_heading("BAB VIII. DAFTAR PUSTAKA", level=1)
    h.runs[0].font.size = Pt(14)
    add_para("1. Russell, S. & Norvig, P. (2021). Artificial Intelligence: A Modern Approach. Fourth Edition. Pearson Education.")
    add_para("2. McCarthy, J. (2007). What is Artificial Intelligence?. Stanford University.")
    add_para("3. World Economic Forum. (2023). The Future of Jobs Report 2023.")
    add_para("4. Kaplan, A. & Haenlein, M. (2019). Siri, in my hand. Business Horizons, 62(1), 15-25.")
    add_para("5. Brynjolfsson, E. & McAfee, A. (2017). Machine, Platform, Crowd. W.W. Norton & Company.")
    add_para("6. UNESCO. (2021). Recommendation on the Ethics of Artificial Intelligence.")
    
    # BAB IX
    h = doc.add_heading("BAB IX. LAMPIRAN", level=1)
    h.runs[0].font.size = Pt(14)
    add_para("Lampiran A: Gambar Ilustrasi AI")
    add_para("[Screenshot berbagai implementasi AI]")
    add_para("")
    add_para("Lampiran B: Timeline Perkembangan AI")
    add_para("- 1956: Konsep AI diperkenalkan di Dartmouth Conference")
    add_para("- 1966: ELIZA, chatbot pertama")
    add_para("- 1997: Deep Blue kalahkan Grandmaster Catur")
    add_para("- 2011: Watson menang di Jeopardy")
    add_para("- 2016: AlphaGo kalahkan Champion Go")
    add_para("- 2020: GPT-3 released")
    add_para("- 2023: ChatGPT launched")
    add_para("- 2024: GPT-4 dan model AI baru")
    add_para("- 2026: AI terintegrasi dalam kehidupan")
    add_para("")
    add_para("Lampiran C: Glossary")
    add_para("- AI: Artificial Intelligence")
    add_para("- ML: Machine Learning")
    add_para("- DL: Deep Learning")
    add_para("- NLP: Natural Language Processing")
    add_para("- NN: Neural Network")
    add_para("- AGI: Artificial General Intelligence")
    
    output = "/root/.nullclaw/workspace/tasks/Laporan_Lengkap_AI.docx"
    doc.save(output)
    return output

if __name__ == "__main__":
    out = buat_laporan()
    print(f"Generated: {out}")
    send_file(out)
    print("Dikirim ke Telegram")
