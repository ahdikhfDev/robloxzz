import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_run_font(run, font_name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Fallback untuk font non-Latin di docx
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def add_para(doc, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p

def add_heading_custom(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        set_run_font(run, font_name='Times New Roman', size=14 if level==1 else 12, bold=True)
    return h

def main():
    doc = Document()
    # Setup default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # COVER
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAPORAN AKADEMIK")
    set_run_font(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTERNET OF THINGS (IoT):\nKONSEP, ARSITEKTUR, IMPLEMENTASI, DAN TANTANGAN MASA DEPAN")
    set_run_font(run, size=14, bold=True)

    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Disusun untuk memenuhi standar akademik perguruan tinggi\nberdasarkan tinjauan jurnal ilmiah terkini")
    set_run_font(run, size=12, italic=True)

    doc.add_page_break()

    # BAB I
    add_heading_custom(doc, "BAB I: PENDAHULUAN", level=1)
    add_heading_custom(doc, "1.1 Latar Belakang", level=2)
    add_para(doc, "Perkembangan teknologi informasi dan komunikasi telah membawa perubahan fundamental dalam cara manusia berinteraksi dengan lingkungan fisik. Internet of Things (IoT) muncul sebagai paradigma revolusioner yang menghubungkan objek-objek fisik ke jaringan internet, memungkinkan pertukaran data secara real-time tanpa intervensi manusia secara langsung. Menurut Gartner (2023), jumlah perangkat IoT yang terhubung secara global diproyeksikan melampaui 25 miliar pada tahun 2025, menandai era baru dalam otomatisasi industri, smart city, dan kehidupan sehari-hari.")
    add_para(doc, "Integrasi sensor, aktuator, mikrokontroler, dan protokol komunikasi nirkabel telah menurunkan biaya implementasi secara signifikan. Namun, adopsi IoT yang masif juga membawa tantangan kompleks terkait keamanan siber, interoperabilitas, skalabilitas, dan privasi data. Oleh karena itu, pemahaman mendalam mengenai arsitektur, standar protokol, dan kerangka kerja keamanan menjadi krusial bagi pengembang, peneliti, dan pembuat kebijakan.")

    add_heading_custom(doc, "1.2 Rumusan Masalah", level=2)
    add_para(doc, "Berdasarkan latar belakang tersebut, rumusan masalah dalam laporan ini adalah:\n1. Bagaimana arsitektur dan komponen fundamental dalam sistem IoT?\n2. Protokol komunikasi apa saja yang dominan digunakan dalam ekosistem IoT?\n3. Bagaimana implementasi IoT pada berbagai sektor industri dan kehidupan?\n4. Apa saja tantangan keamanan dan privasi yang dihadapi serta solusi mitigasinya?")

    add_heading_custom(doc, "1.3 Tujuan dan Manfaat", level=2)
    add_para(doc, "Laporan ini bertujuan untuk memberikan tinjauan komprehensif mengenai konsep IoT, mulai dari lapisan arsitektur hingga implementasi praktis. Manfaat yang diharapkan adalah sebagai referensi akademik bagi mahasiswa teknik informatika, sistem informasi, dan elektro, serta sebagai panduan awal bagi praktisi yang ingin mengadopsi solusi IoT dalam proyek mereka.")

    doc.add_page_break()

    # BAB II
    add_heading_custom(doc, "BAB II: TINJAUAN PUSTAKA", level=1)
    add_heading_custom(doc, "2.1 Definisi dan Konsep Dasar IoT", level=2)
    add_para(doc, "Internet of Things didefinisikan sebagai jaringan objek fisik yang tertanam dengan sensor, perangkat lunak, dan teknologi lainnya untuk tujuan menghubungkan dan bertukar data dengan perangkat dan sistem lain melalui internet. Konsep ini pertama kali dipopulerkan oleh Kevin Ashton pada tahun 1999, yang menekankan pada kemampuan komputer untuk mengumpulkan dan memahami data dunia nyata tanpa bergantung pada input manusia.")

    add_heading_custom(doc, "2.2 Arsitektur IoT", level=2)
    add_para(doc, "Arsitektur IoT umumnya dibagi menjadi empat lapisan utama:\n1. Perception Layer (Lapisan Persepsi): Terdiri dari sensor, aktuator, dan perangkat edge yang bertugas mengumpulkan data fisik (suhu, kelembaban, gerakan, dll).\n2. Network Layer (Lapisan Jaringan): Bertanggung jawab mentransmisikan data dari perangkat edge ke cloud atau server menggunakan protokol seperti Wi-Fi, Bluetooth, Zigbee, LoRaWAN, atau seluler (4G/5G).\n3. Middleware/Processing Layer: Memproses, menyimpan, dan menganalisis data. Seringkali melibatkan edge computing atau cloud computing.\n4. Application Layer: Antarmuka pengguna dan layanan spesifik seperti smart home dashboard, monitoring industri, atau sistem kesehatan digital.")

    add_heading_custom(doc, "2.3 Protokol Komunikasi IoT", level=2)
    add_para(doc, "Pemilihan protokol sangat bergantung pada kebutuhan daya, jangkauan, dan bandwidth. Protokol utama meliputi:\n- MQTT (Message Queuing Telemetry Transport): Ringan, berbasis publish-subscribe, ideal untuk jaringan tidak stabil.\n- CoAP (Constrained Application Protocol): Dirancang untuk perangkat dengan sumber daya terbatas, menggunakan UDP.\n- HTTP/REST: Standar web, mudah diintegrasikan namun lebih berat.\n- LoRaWAN & NB-IoT: Untuk komunikasi jarak jauh dengan konsumsi daya rendah (LPWAN).")

    doc.add_page_break()

    # BAB III
    add_heading_custom(doc, "BAB III: METODOLOGI & IMPLEMENTASI", level=1)
    add_heading_custom(doc, "3.1 Perangkat Keras (Hardware)", level=2)
    add_para(doc, "Implementasi IoT memerlukan mikrokontroler atau single-board computer sebagai otak sistem. Platform populer meliputi ESP32/ESP8266 (terintegrasi Wi-Fi/Bluetooth), Arduino (untuk prototyping cepat), dan Raspberry Pi (untuk pemrosesan edge yang lebih kompleks). Sensor yang umum digunakan meliputi DHT22 (suhu/kelembaban), PIR (gerakan), dan MQ-series (gas).")

    add_heading_custom(doc, "3.2 Perangkat Lunak & Platform Cloud", level=2)
    add_para(doc, "Sisi perangkat lunak melibatkan firmware (biasanya C/C++ atau MicroPython), broker MQTT (seperti Mosquitto atau HiveMQ), dan platform cloud (AWS IoT Core, Google Cloud IoT, atau Blynk). Analisis data dapat dilakukan menggunakan stack ELK (Elasticsearch, Logstash, Kibana) atau tools machine learning seperti TensorFlow Lite untuk prediksi di edge.")

    add_heading_custom(doc, "3.3 Alur Kerja Sistem", level=2)
    add_para(doc, "Alur kerja standar IoT dimulai dari akuisisi data oleh sensor, konversi sinyal analog ke digital, transmisi via protokol nirkabel ke gateway/broker, pemrosesan di cloud/edge, penyimpanan di database time-series, dan visualisasi melalui dashboard web/mobile. Sistem juga dapat dilengkapi dengan rule engine untuk triggering otomatis (misal: kirim notifikasi jika suhu > 40°C).")

    doc.add_page_break()

    # BAB IV
    add_heading_custom(doc, "BAB IV: STUDI KASUS & ANALISIS", level=1)
    add_heading_custom(doc, "4.1 Smart Agriculture", level=2)
    add_para(doc, "IoT telah merevolusi pertanian presisi. Sensor tanah memantau kelembaban, pH, dan nutrisi secara real-time. Data dikirim ke cloud untuk analisis irigasi otomatis dan pemupukan tepat dosis. Studi oleh FAO (2022) menunjukkan peningkatan hasil panen hingga 20% dan penghematan air hingga 30% dengan adopsi sistem IoT berbasis LoRaWAN di lahan pertanian skala menengah.")

    add_heading_custom(doc, "4.2 Smart Home & Building Automation", level=2)
    add_para(doc, "Otomasi rumah menggunakan IoT mencakup pencahayaan adaptif, HVAC cerdas, keamanan berbasis kamera AI, dan manajemen energi. Protokol Zigbee dan Matter semakin diadopsi untuk memastikan interoperabilitas antar merek. Tantangan utama adalah fragmentasi ekosistem dan kerentanan terhadap serangan DDoS jika perangkat tidak di-patch secara berkala.")

    add_heading_custom(doc, "4.3 Industrial IoT (IIoT) & Industry 4.0", level=2)
    add_para(doc, "Di sektor manufaktur, IIoT memungkinkan predictive maintenance, digital twin, dan optimasi rantai pasok. Sensor getaran dan termal pada mesin mendeteksi anomali sebelum kegagalan terjadi, mengurangi downtime hingga 40%. Integrasi dengan AI dan big data analytics menjadi kunci transformasi menuju pabrik otonom.")

    doc.add_page_break()

    # BAB V
    add_heading_custom(doc, "BAB V: PENUTUP", level=1)
    add_heading_custom(doc, "5.1 Kesimpulan", level=2)
    add_para(doc, "Internet of Things telah berkembang dari konsep eksperimental menjadi tulang punggung transformasi digital di berbagai sektor. Arsitektur berlapis, protokol komunikasi yang efisien, dan integrasi cloud/edge computing memungkinkan skalabilitas dan real-time processing. Namun, keberhasilan adopsi IoT sangat bergantung pada penanganan isu keamanan, privasi data, standarisasi protokol, dan kesiapan infrastruktur jaringan.")

    add_heading_custom(doc, "5.2 Saran", level=2)
    add_para(doc, "Untuk pengembangan lebih lanjut, disarankan:\n1. Menerapkan security-by-design pada setiap lapisan arsitektur IoT.\n2. Mengadopsi standar terbuka seperti Matter untuk interoperabilitas.\n3. Meningkatkan literasi keamanan siber bagi pengembang dan pengguna akhir.\n4. Melakukan riset lebih lanjut pada edge AI untuk mengurangi latensi dan beban cloud.\n5. Kolaborasi antara akademisi, industri, dan regulator untuk menyusun kerangka kebijakan IoT yang adaptif.")

    doc.add_page_break()

    # DAFTAR PUSTAKA
    add_heading_custom(doc, "DAFTAR PUSTAKA", level=1)
    refs = [
        "Ashton, K. (2009). That 'Internet of Things' Thing. RFID Journal. https://www.rfidjournal.com/articles/view?4986",
        "Al-Fuqaha, A., Guizani, M., Mohammadi, M., Aledhari, M., & Ayyash, M. (2015). Internet of Things: A Survey on Enabling Technologies, Protocols, and Applications. IEEE Communications Surveys & Tutorials, 17(4), 2347-2376. https://doi.org/10.1109/COMST.2015.2444095",
        "Gubbi, J., Buyya, R., Marusic, S., & Palaniswami, M. (2013). Internet of Things (IoT): A vision, architectural elements, and future directions. Future Generation Computer Systems, 29(7), 1645-1660. https://doi.org/10.1016/j.future.2013.01.010",
        "Sisinni, E., Saifullah, A., Han, S., Jennehag, U., & Gidlund, M. (2018). Industrial Internet of Things: Challenges, Opportunities, and Directions. IEEE Transactions on Industrial Informatics, 14(11), 4724-4734. https://doi.org/10.1109/TII.2018.2852491",
        "Ray, P. P. (2018). A survey on Internet of Things architectures. Journal of King Saud University - Computer and Information Sciences, 30(3), 291-319. https://doi.org/10.1016/j.jksuci.2016.10.003",
        "FAO. (2022). The State of Food and Agriculture 2022: Leveraging automation in agriculture for transforming agrifood systems. Food and Agriculture Organization of the United Nations. https://www.fao.org/3/cc2220en/cc2220en.pdf",
        "Matter Specification. (2023). Connectivity Standards Alliance. https://csa-iot.org/all-solutions/matter/"
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"[{i}] {ref}")
        set_run_font(run, size=11)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "Laporan_IoT_Akademik.docx")
    doc.save(output_path)
    print(f"✅ Laporan berhasil dibuat: {output_path}")

if __name__ == "__main__":
    main()
